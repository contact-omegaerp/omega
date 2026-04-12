"""
gen_handover1.py — Tạo OMEGA-WEB-HANDOVER1.docx
Nhặt nội dung từ file hướng dẫn deploy cũ và merge vào HANDOVER.docx
"""
import shutil, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN_DARK  = RGBColor(0x0d, 0x5c, 0x38)
GREEN_MED   = RGBColor(0x1a, 0x7a, 0x50)
WHITE       = RGBColor(0xff, 0xff, 0xff)

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '%02X%02X%02X' % rgb)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = GREEN_DARK; r.font.size = Pt(15)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = GREEN_MED; r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs: r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)

def body(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.bold = bold
    p.paragraph_format.space_after = Pt(4)

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F0F4F8')
    pPr.append(shd)

def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, (0x0d, 0x5c, 0x38))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_bg(cell, (0xe8, 0xf5, 0xe9))
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()

# ── Mở file gốc ───────────────────────────────────────────────────────────────
src = 'OMEGA-WEB-HANDOVER.docx'
dst = 'OMEGA-WEB-HANDOVER1.docx'
shutil.copy(src, dst)
doc = Document(dst)

# ──────────────────────────────────────────────────────────────────────────────
# MỤC 4B: Hướng dẫn deploy GitHub Pages chi tiết
# ──────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, 'MUC 4B. HUONG DAN DEPLOY GITHUB PAGES CHI TIET')
body(doc, 'Quy trinh day du tung buoc de cap nhat va trien khai website len omega.com.vn.')

h2(doc, 'Phan 1 - Day code len GitHub va kich hoat GitHub Pages')

h3(doc, '1.1 Chuan bi repository')
bullet(doc, 'Repository hien tai: github.com/lazinet/omega (da co san - khong can tao lai)')
bullet(doc, 'Che do: Public (GitHub Pages mien phi yeu cau public repo)')
bullet(doc, 'File trang chu phai la index.html o thu muc goc (root) cua repo')

h3(doc, '1.2 Push code tu VS Code - Quy trinh hang ngay')
body(doc, 'Mo Terminal trong VS Code (Ctrl + `) va chay theo thu tu:')
code(doc, 'Step 1: cd D:\\Dev_SW\\projects\\omega')
code(doc, 'Step 2: python -X utf8 auto-omega/run_all.py   <- chay pipeline truoc')
code(doc, 'Step 3: git add .')
code(doc, 'Step 4: git commit -m "Mo ta thay doi ngan gon"')
code(doc, 'Step 5: git push')
bullet(doc, 'GitHub Pages deploy tu dong trong < 60 giay sau git push')
bullet(doc, 'Kiem tra tien trinh deploy: github.com/lazinet/omega -> tab Actions')

h3(doc, '1.3 Kich hoat GitHub Pages (chi lam 1 lan, da hoan thanh)')
bullet(doc, 'Settings -> Pages -> Source: Deploy from a branch')
bullet(doc, 'Branch: main | Folder: /(root) -> Click Save')
bullet(doc, 'Site live tai: https://omega.com.vn (sau khi gan domain)')

h2(doc, 'Phan 2 - Cau hinh domain rieng (omega.com.vn)')

h3(doc, '2.1 Custom domain tren GitHub')
bullet(doc, 'Settings -> Pages -> Custom domain -> go: omega.com.vn -> Save')
bullet(doc, 'GitHub tu tao file CNAME trong repo - khong can tao thu cong')

h3(doc, '2.2 Cau hinh DNS records')
body(doc, 'Tai nha cung cap domain (PA Vietnam), tao cac DNS records sau:')
table(doc,
    ['Loai record', 'Name/Host', 'Value / Points to', 'Muc dich'],
    [
        ['A', '@', '185.199.108.153', 'GitHub Pages IPv4 #1'],
        ['A', '@', '185.199.109.153', 'GitHub Pages IPv4 #2'],
        ['A', '@', '185.199.110.153', 'GitHub Pages IPv4 #3'],
        ['A', '@', '185.199.111.153', 'GitHub Pages IPv4 #4'],
        ['CNAME', 'www', 'lazinet.github.io.', 'Subdomain www -> GitHub'],
    ],
    col_widths=[3, 2.5, 5.5, 4.5]
)
bullet(doc, 'DNS propagation: 5 phut den 48 gio (thuong 30 phut - 2 gio)')
bullet(doc, 'Bat Enforce HTTPS sau khi DNS tro dung (Settings -> Pages -> Enforce HTTPS)')

h3(doc, '2.3 Ke hoach Cloudflare CDN (sap trien khai)')
bullet(doc, 'Parking domain qua Cloudflare thay vi tro thang vao GitHub Pages')
bullet(doc, 'Loi ich: CDN co PoP Ha Noi + TP.HCM -> giam latency ~40-60% cho nguoi dung Viet Nam')
bullet(doc, 'Tinh nang them: Page Rules de 301 redirect URL cu WordPress -> URL moi Static')
bullet(doc, 'Buoc thuc hien: doi Nameserver sang Cloudflare NS -> Proxy (mau cam) thay vi DNS only (mau xam)')
bullet(doc, 'Ho tro: Lien he ky thuat de duoc ho tro cau hinh khi can')

# ──────────────────────────────────────────────────────────────────────────────
# MỤC 5B: Analytics nâng cao
# ──────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, 'MUC 5B. CHIEN LUOC ANALYTICS & MARKETING NANG CAO')

h2(doc, '1. Ba tang Analytics - Lo trinh tong the')
body(doc, 'Hien tai OMEGA dang o Tang 1. Muc tieu tien len Tang 3 trong Q2/2026.')
table(doc,
    ['Tang', 'Cong cu', 'Cau hoi tra loi', 'Trang thai'],
    [
        ['Tang 1 - Do dac', 'GA4 + Clarity', 'Ai den, tu dau, xem gi, bao lau; click dau, scroll den dau', 'Hoan thanh'],
        ['Tang 2 - Hieu', 'GA4 reports + Clarity heatmap', 'Trang nao bounce cao? Nut CTA nao bi bo qua?', 'Cho data (4-6 tuan)'],
        ['Tang 3 - Hanh dong', 'GA4 Conversion + A/B test', 'Trang nao mang ve lead? Source nao convert tot?', 'Ke hoach Q2/2026'],
    ],
    col_widths=[3.5, 4, 6, 3]
)

h2(doc, '2. Conversion Goals can thiet lap trong GA4')
body(doc, 'Hien GA4 chi dem pageview - chua biet ai "co y dinh mua". Can danh dau cac event:')
table(doc,
    ['Event', 'Trigger', 'Ten Conversion'],
    [
        ['click_lien_he', 'Click vao lien-he.html hoac nut "Tu van mien phi"', 'Lead Intent'],
        ['click_phone', 'Click tel: hoac zalo.me trong floating contact', 'Phone / Zalo Click'],
        ['form_submit', 'Gui form lien he thanh cong', 'Lead Submitted'],
        ['click_case_study', 'Click xem chi tiet khach hang tieu bieu', 'Engaged User'],
    ],
    col_widths=[4, 7.5, 4]
)
body(doc, 'Sau khi cai dat: biet duoc trang nao, traffic source nao, nganh nao dang mang ve khach hang tiem nang.')

h2(doc, '3. Google Search Console - Hanh dong tiep theo')
table(doc,
    ['Buoc', 'Hanh dong', 'Cong cu'],
    [
        ['1', 'Kiem tra tai khoan Google cu co property omega.com.vn chua', 'Lien he ben cu'],
        ['2', 'Neu co: yeu cau Transfer Ownership -> Manage Property -> Users', 'GSC Dashboard'],
        ['3', 'Neu chua: tao Domain property moi tai search.google.com/search-console', 'GSC Dashboard'],
        ['4', 'Xac minh qua DNS TXT record hoac qua GA4 (cung tai khoan)', 'DNS Manager / GA4'],
        ['5', 'Submit sitemap: https://omega.com.vn/sitemap.xml', 'GSC -> Sitemaps'],
        ['6', 'Ket noi GSC <-> GA4: GSC -> Settings -> Associations -> Link to GA4', 'GSC + GA4'],
    ],
    col_widths=[1.2, 9, 4]
)

h2(doc, '4. AEO - Answer Engine Optimization (dai han, 0 dong)')
body(doc, 'Toi uu noi dung de AI (ChatGPT, Gemini, Perplexity, Copilot) tu cite OMEGA khi tra loi cau hoi ve ERP.')
bullet(doc, 'Viet bai authoritative: "ERP la gi", "So sanh phan mem ERP Viet Nam", "Cach chon ERP cho nganh nhua"')
bullet(doc, 'Noi dung phai co du lieu that: so lieu, case study, kinh nghiem 16 nam trien khai')
bullet(doc, 'Submit sitemap len Bing Webmaster Tools (Microsoft Copilot dung Bing de search)')
bullet(doc, 'Dam bao robots.txt cho phep AI crawlers truy cap')
bullet(doc, 'Backlink tu cac trang tin tuc cong nghe, bao kinh te Viet Nam')

h2(doc, '5. Ke hoach hanh dong cu the')
table(doc,
    ['Uu tien', 'Hanh dong', 'Chi phi', 'Timeline'],
    [
        ['1 - Ngay lap tuc', 'Submit sitemap len GSC + Bing Webmaster Tools', '0 dong', 'Tuan 1 (thang 4)'],
        ['1 - Ngay lap tuc', 'Thiet lap Conversion Goals trong GA4', '0 dong', 'Tuan 1 (thang 4)'],
        ['2 - Ngan han', 'Ket noi Clarity <-> GA4 trong dashboard Clarity', '0 dong', 'Tuan 2 (thang 4)'],
        ['3 - Trung han', 'Google Ads tu khoa ERP + nganh muc tieu', 'Co phi', 'Q2/2026'],
        ['3 - Trung han', 'Microsoft Ads (Bing) - import tu Google Ads', 'Co phi, re hon 30-50%', 'Q2/2026'],
        ['4 - Dai han', 'Chuoi bai AEO: Q&A chuyen sau tung nganh', '0 dong (content)', 'Q3/2026'],
    ],
    col_widths=[3.5, 5.5, 2.5, 3]
)

# ──────────────────────────────────────────────────────────────────────────────
# MỤC 6: Lộ trình hệ thống mở rộng
# ──────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, 'MUC 6. LO TRINH HE THONG MO RONG')

h2(doc, 'Tong quan 4 he thong ke hoach')
table(doc,
    ['#', 'He thong', 'Quyet dinh', 'Chi phi', 'Do phuc tap'],
    [
        ['1', 'Online Support Status', 'Lam ngay - Google Sheets backend', 'Mien phi', 'Thap'],
        ['2', 'AI CV Scoring / Mini ATS', 'Hoan - da co base tuong tu', 'API cost nho', 'Trung binh'],
        ['3', 'Email CRM (Brevo)', 'Nghien cuu - chua trien khai', 'Mien phi (free tier)', 'Thap-Trung'],
        ['4', 'Cong Affiliate / CTV', 'Sau cung - validate truoc', 'Bien doi', 'Cao'],
    ],
    col_widths=[0.7, 4.3, 4.5, 2.5, 2.5]
)

h2(doc, '1. Online Support Status')
body(doc, 'Hien thi trang thai team tu van ngay tren website, cap nhat qua Google Sheets.')
bullet(doc, 'Backend: Google Sheets co cot: team, status (Online/Offline/Ban), message, next_available')
bullet(doc, 'Publish as CSV -> Widget JS tren site fetch moi 5 phut -> hien thi badge mau trong section Dich vu')
bullet(doc, 'Admin tu cap nhat Sheet -> site tu doi (khong can deploy, khong can coder)')
bullet(doc, 'Mo rong: tich hop Google Calendar -> tu detect le tet -> "Nghi le" tu dong')

h2(doc, '2. Email CRM voi Brevo')
body(doc, 'Nen tang email marketing + CRM mien phi, tich hop truc tiep voi form tren site.')
h3(doc, 'Cach hoat dong')
bullet(doc, 'Form HTML Omega -> Brevo API -> Contact tu dong vao List tuong ung')
bullet(doc, 'Free tier: 300 email/ngay, unlimited contacts, 9 automation workflow')
h3(doc, 'Phan loai Lists')
table(doc,
    ['List', 'Nguon', 'Muc dich'],
    [
        ['Lead moi', 'Form "Nhan tu van" tren site', 'Nurture -> convert thanh khach hang'],
        ['Khach hang', 'Import tu du lieu noi bo', 'Cham soc, upsell, gia han'],
        ['Newsletter', 'Form cuoi trang', 'Tin tuc, san pham moi, bai viet'],
        ['CTV/Doi tac', 'Form dang ky CTV (tuong lai)', 'Thong tin hop tac, hoa hong'],
    ],
    col_widths=[3, 5, 6.5]
)
h3(doc, 'Automation workflows')
bullet(doc, 'Lead moi: xac nhan ngay -> case study phu hop nganh (D+1) -> follow-up (D+3) -> nurture hang thang')
bullet(doc, 'Khach hang cu: chuc mung le tet tu dong (30/4, 2/9, Tet), nhac gia han hop dong')
bullet(doc, 'Custom fields: COMPANY, INDUSTRY, SOURCE, STATUS -> personalize theo tung doi tuong')
h3(doc, 'Can xac dinh truoc khi trien khai')
bullet(doc, 'Danh sach khach hang hien tai luu o dau? (Excel, phan mem noi bo?)')
bullet(doc, 'Form lien-he.html hien co ai nhan email khong?')
bullet(doc, 'Uu tien cham soc khach cu hay nurture lead moi truoc?')

h2(doc, '3. AI CV Scoring / Mini ATS')
body(doc, 'Form upload CV -> AI phan tich theo JD -> shortlist ung vien tu dong, giam tai cho HR.')
bullet(doc, 'Luu y ky thuat: khong the goi AI API truc tiep tu HTML (lo API key)')
bullet(doc, 'Can mot lop proxy: Cloudflare Worker (mien phi 100k request/ngay) -> an key, forward den Claude API')
body(doc, 'Flow du an:')
code(doc, 'Form upload CV (PDF/DOCX) -> Cloudflare Worker -> Claude API')
code(doc, '-> Tra ve diem + nhan xet -> Luu vao Airtable -> Email HR tu dong')

h2(doc, '4. Cong Affiliate / Cong tac vien')
h3(doc, 'Phien ban don gian (khuyen nghi bat dau)')
bullet(doc, 'Form dang ky CTV -> Airtable -> Email xac nhan tu dong')
bullet(doc, 'Link tracking: UTM parameter (GA4 da co san, khong can them gi)')
bullet(doc, 'Bao cao hoa hong: export thu cong tu GA4 + CRM | Chi phi: 0 dong')
h3(doc, 'Phien ban day du (neu can portal that)')
bullet(doc, 'Dung Memberstack hoac Outseta (co free tier) cho auth + dashboard')
bullet(doc, 'Hoac: Supabase (auth + DB free) + Netlify Functions')
bullet(doc, 'Uoc tinh: 2-4 tuan dev nghiem tuc')

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(dst)
size = os.path.getsize(dst)
print(f'Da luu: {dst}')
print(f'Kich thuoc: {size/1024:.1f} KB')

if __name__ == '__main__':
    pass
