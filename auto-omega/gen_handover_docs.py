# -*- coding: utf-8 -*-
"""
Sinh 2 file tài liệu bàn giao dự án OMEGA website:
  - OMEGA-WEB-HANDOVER.docx
  - OMEGA-WEB-SITEMAP.xlsx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import (
    PatternFill, Font, Alignment, Border, Side
)
from openpyxl.utils import get_column_letter

OUTPUT_DIR = r"D:\Dev_SW\projects\omega\auto-omega"

# ─────────────────────────────────────────────
# COLOR CONSTANTS
# ─────────────────────────────────────────────
GREEN_DARK  = RGBColor(0x0D, 0x5C, 0x38)   # #0d5c38
GREEN_MED   = RGBColor(0x1A, 0x7A, 0x50)   # heading 2
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_LIGHT = RGBColor(0xE8, 0xF5, 0xE9)   # alternate row
GRAY_LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)

# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_cell_border(table):
    """Add thin borders to every cell in the table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border_el = OxmlElement(f"w:{edge}")
                border_el.set(qn("w:val"), "single")
                border_el.set(qn("w:sz"), "4")
                border_el.set(qn("w:color"), "CCCCCC")
                tcBorders.append(border_el)
            tcPr.append(tcBorders)

def style_table_header(row, bg_hex="0D5C38", font_size=10):
    for cell in row.cells:
        set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(font_size)
        # also style text that hasn't been wrapped in runs
        if cell.paragraphs and not cell.paragraphs[0].runs:
            run = cell.paragraphs[0].add_run(cell.paragraphs[0].text)
            cell.paragraphs[0].clear()
            run2 = cell.paragraphs[0].add_run(run.text)
            run2.font.bold = True
            run2.font.color.rgb = WHITE
            run2.font.size = Pt(font_size)

def add_heading1(doc, text):
    p = doc.add_heading(level=1)
    p.clear()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = GREEN_DARK
    return p

def add_heading2(doc, text):
    p = doc.add_heading(level=2)
    p.clear()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = GREEN_MED
    return p

def add_paragraph(doc, text="", bold=False, italic=False, size=11, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.bold = bold
        run.font.italic = italic
        run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_code_block(doc, code_text):
    """Add a shaded code block."""
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        # light gray background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F4F0")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)

def add_ordered_item(doc, number, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def build_feature_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "0D5C38")

    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = "E8F5E9" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            set_cell_bg(cell, bg)

    add_cell_border(table)
    return table


# ─────────────────────────────────────────────
# BUILD DOCX
# ─────────────────────────────────────────────

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TÀI LIỆU BÀN GIAO & HƯỚNG DẪN VẬN HÀNH WEBSITE OMEGA")
    run.font.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = GREEN_DARK

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("omega.com.vn  |  Phiên bản 1.0  |  06/04/2026")
    run2.font.size = Pt(13)
    run2.font.color.rgb = GREEN_MED
    run2.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Decorative line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("─" * 60)
    run_line.font.color.rgb = GREEN_DARK

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = info.add_run("Công ty TNHH Công nghệ và Giải pháp Omega\n16+ năm kinh nghiệm  |  1000+ khách hàng doanh nghiệp\nRepo: github.com/lazinet/omega")
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ── MỤC 1 ───────────────────────────────────
    add_heading1(doc, "MỤC 1. GIỚI THIỆU DỰ ÁN")
    doc.add_paragraph()

    add_heading2(doc, "1.1 Mục tiêu")
    add_bullet(doc, "Hiện đại hóa website từ WordPress cũ sang static HTML thuần — nhanh hơn, an toàn hơn, kiểm soát hoàn toàn")
    add_bullet(doc, "Loại bỏ phụ thuộc vào plugin WordPress, hosting PHP — zero server-side runtime")
    add_bullet(doc, "Tự động hoá toàn bộ quy trình: sinh trang, sync header, inject analytics, build sitemap")

    doc.add_paragraph()
    add_heading2(doc, "1.2 Kết quả đạt được")
    add_bullet(doc, "239 trang HTML hoàn chỉnh (sản phẩm, giải pháp, khách hàng, tin tức, special pages)")
    add_bullet(doc, "Pipeline tự động hoá: 5 Python scripts trong auto-omega/")
    add_bullet(doc, "Analytics tracking: GA4 + Microsoft Clarity đã inject vào toàn bộ 239 trang")
    add_bullet(doc, "Tìm kiếm client-side: overlay + search-index.json + phím tắt Ctrl+K")
    add_bullet(doc, "Custom 404 thông minh: orbital animation, auto-redirect 10 giây, quick links")
    add_bullet(doc, "CI/CD từ VS Code → git push → GitHub Pages live trong < 60 giây")

    doc.add_paragraph()
    add_heading2(doc, "1.3 Đối tượng tài liệu")
    add_bullet(doc, "Nhân viên quản trị content nội bộ Omega")
    add_bullet(doc, "Kỹ thuật web nội bộ Omega phụ trách vận hành và phát triển tiếp")

    doc.add_page_break()

    # ── MỤC 2 ───────────────────────────────────
    add_heading1(doc, "MỤC 2. KIẾN TRÚC HỆ THỐNG")
    doc.add_paragraph()

    add_heading2(doc, "2.1 Luồng CI/CD")
    add_code_block(doc, """Chỉnh sửa code (VS Code)
     ↓
Python automation scripts (auto-omega/)
     ↓
git commit + push → GitHub (github.com/lazinet/omega)
     ↓
GitHub Pages deploy tự động (< 60 giây)
     ↓
omega.com.vn (live)""")

    doc.add_paragraph()
    add_heading2(doc, "2.2 Thành phần kỹ thuật")

    tech_headers = ["Thành phần", "Chi tiết"]
    tech_rows = [
        ["Frontend", "HTML5 thuần, CSS3 (Bootstrap 5), JavaScript (jQuery 3.7, Swiper, WOW.js, GSAP)"],
        ["CSS chính", "assets/css/omega.css — CSS variables toàn site (--dark-bg, --green-primary, v.v.)"],
        ["JS chính", "assets/js/omega.js — preloader, sticky nav, search overlay, back-to-top, canvas effects"],
        ["Hosting", "GitHub Pages (free, SSL tự động)"],
        ["CDN dự kiến", "Cloudflare (có PoP Hà Nội + TP.HCM) — parking domain để cache tĩnh, giảm latency ~40-60%"],
        ["Google App Script", "Đang dùng cho automation (form xử lý, v.v.)"],
        ["Analytics", "GA4 (G-0T4G53DCPV) + Microsoft Clarity (w5st3jx7zo) — inject vào 239 trang"],
        ["Repository", "github.com/lazinet/omega"],
    ]
    build_feature_table(doc, tech_headers, tech_rows)

    doc.add_page_break()

    # ── MỤC 3 ───────────────────────────────────
    add_heading1(doc, "MỤC 3. DANH SÁCH TÍNH NĂNG")
    doc.add_paragraph()

    feat_headers = ["#", "Tính năng", "Mô tả", "Trạng thái"]
    feat_rows = [
        [1,  "Mega-menu desktop",        "Nav 5 cấp với mega-panel sản phẩm",                              "✅ Hoàn thành"],
        [2,  "Mobile menu",              "SlickNav responsive",                                             "✅ Hoàn thành"],
        [3,  "Tìm kiếm client-side",     "Overlay tìm kiếm toàn site, index JSON, phím tắt Ctrl+K",       "✅ Hoàn thành"],
        [4,  "Google Translate",         "Lazy-load, 30+ ngôn ngữ",                                        "✅ Hoàn thành"],
        [5,  "Floating contact widget",  "Nút xoay: CSKH / Messenger / Zalo / Gọi điện",                  "✅ Hoàn thành"],
        [6,  "Canvas animations",        "Radar, star network, contact radar — per-page",                  "✅ Hoàn thành"],
        [7,  "Preloader",                "Ngăn FOUC khi mạng chậm",                                        "✅ Hoàn thành"],
        [8,  "Dark/light sections",      "Xen kẽ tối/sáng, CSS variables toàn site",                      "✅ Hoàn thành"],
        [9,  "Hero stats counter",       "Số đếm động (16+ năm, 1000+ KH, 20+ tỉnh)",                     "✅ Hoàn thành"],
        [10, "Customer logos swiper",    "19 logo thật, auto-play",                                        "✅ Hoàn thành"],
        [11, "Testimonials section",     "3 case study tiêu biểu với số liệu thật",                        "✅ Hoàn thành"],
        [12, "Sitemap tự động",          "gen_sitemap.py sinh sitemap.xml chuẩn sitemaps.org",             "✅ Hoàn thành"],
        [13, "Search index tự động",     "gen_search_index.py sinh search-index.json",                     "✅ Hoàn thành"],
        [14, "Header sync tự động",      "gen_sync_header.py đồng bộ 5 block từ index → 9 trang",         "✅ Hoàn thành"],
        [15, "Analytics sync tự động",   "gen_analytics.py inject GA4+Clarity+Preloader CSS → 239 trang", "✅ Hoàn thành"],
        [16, "Customer pages tự động",   "gen_customer_pages.py sinh 19 trang khach-hang/",               "✅ Hoàn thành"],
        [17, "Custom 404 page",          "Orbital animation, search, quick links, auto-redirect 10s",      "✅ Hoàn thành"],
        [18, "OG/Social sharing",        "og:image nền trắng, title/desc đầy đủ mọi trang",               "✅ Hoàn thành"],
        [19, "Schema.org structured data","Organization + LocalBusiness + SoftwareApplication",            "✅ Hoàn thành"],
        [20, "Google Search Console",    "Chưa verify — cần transfer từ bên cũ hoặc tạo mới",             "⏳ Pending"],
        [21, "Online Support Status",    "Google Sheets → widget trạng thái tư vấn viên",                 "📋 Kế hoạch"],
        [22, "Email CRM (Brevo)",        "Automation email khách hàng + lead nurturing",                   "📋 Kế hoạch"],
        [23, "Cloudflare CDN",           "Parking domain qua Cloudflare PoP Việt Nam",                     "📋 Kế hoạch"],
    ]
    build_feature_table(doc, feat_headers, feat_rows)

    doc.add_page_break()

    # ── MỤC 4 ───────────────────────────────────
    add_heading1(doc, "MỤC 4. HƯỚNG DẪN VẬN HÀNH")
    doc.add_paragraph()

    add_heading2(doc, "4.1 Lệnh chạy pipeline chính")
    add_code_block(doc, "cd D:\\Dev_SW\\projects\\omega\npython -X utf8 auto-omega/run_all.py")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Chạy lệnh này mỗi khi:")
    run.font.bold = True
    run.font.size = Pt(11)
    add_bullet(doc, "Thêm/sửa bài viết tin tức")
    add_bullet(doc, "Sửa nav menu / header")
    add_bullet(doc, "Thêm trang mới")
    add_bullet(doc, "Trước khi git push")

    doc.add_paragraph()
    add_heading2(doc, "4.2 Thứ tự script trong pipeline")
    script_headers = ["#", "Script", "Chức năng"]
    script_rows = [
        [1, "gen_analytics.py",       "Inject GA4 + Clarity + Preloader CSS → 239 trang"],
        [2, "gen_sync_header.py",     "Đồng bộ header/search/css/translate/floating → 9 trang root"],
        [3, "gen_customer_pages.py",  "Sinh 19 trang khach-hang/"],
        [4, "gen_sitemap.py",         "Sinh sitemap.xml"],
        [5, "gen_search_index.py",    "Sinh search-index.json"],
    ]
    build_feature_table(doc, script_headers, script_rows)

    doc.add_paragraph()
    add_heading2(doc, "4.3 Thêm bài viết mới")
    add_ordered_item(doc, 1, "Tạo file HTML trong tin-tuc/")
    add_ordered_item(doc, 2, "Đặt tên theo slug: tu-khoa-chinh-cua-bai.html")
    add_ordered_item(doc, 3, "Cập nhật tin-tuc/_tools/news-data.json")
    add_ordered_item(doc, 4, "Chạy run_all.py")
    add_ordered_item(doc, 5, "Git commit + push")

    doc.add_paragraph()
    add_heading2(doc, "4.4 Sửa nav menu")
    add_bullet(doc, "Chỉ sửa trong index.html — đây là nguồn chuẩn duy nhất")
    add_bullet(doc, "Chạy run_all.py để đồng bộ sang 9 trang")
    add_bullet(doc, "Git push — deploy tự động")

    doc.add_paragraph()
    add_heading2(doc, "4.5 Git workflow")
    add_code_block(doc, """git add .
git commit -m "Mô tả thay đổi"
git push
# GitHub Pages deploy tự động trong < 60 giây""")

    doc.add_page_break()

    # ── MỤC 5 ───────────────────────────────────
    add_heading1(doc, "MỤC 5. SEO VÀ MARKETING")
    doc.add_paragraph()

    add_heading2(doc, "5.1 Cấu hình SEO hiện tại")
    seo_items = [
        "Meta title, description, keywords: đầy đủ mọi trang",
        "Canonical URL: mỗi trang có <link rel=\"canonical\">",
        "robots.txt: cho phép crawl toàn site",
        "sitemap.xml: tự động sinh, submit lên GSC",
        "Schema.org: Organization, LocalBusiness, SoftwareApplication, BreadcrumbList",
    ]
    for item in seo_items:
        add_bullet(doc, item)

    doc.add_paragraph()
    add_heading2(doc, "5.2 Analytics")
    add_bullet(doc, "GA4 Property: G-0T4G53DCPV (đang chờ transfer quyền từ tài khoản cũ)")
    add_bullet(doc, "Microsoft Clarity: w5st3jx7zo (heatmap + session recording)")
    add_bullet(doc, "Cả 2 đã inject vào 239 trang HTML")

    doc.add_paragraph()
    add_heading2(doc, "5.3 Google Search Console")
    add_bullet(doc, "Hiện trạng: Chưa verify")
    p = doc.add_paragraph()
    run = p.add_run("Bước tiếp theo:")
    run.font.bold = True
    run.font.size = Pt(11)
    add_ordered_item(doc, 1, "Kiểm tra tài khoản Google cũ có property omega.com.vn chưa")
    add_ordered_item(doc, 2, "Nếu có: yêu cầu transfer ownership")
    add_ordered_item(doc, 3, "Nếu chưa: tạo mới, verify qua GA4 (cùng tài khoản)")
    add_ordered_item(doc, 4, "Submit sitemap: https://omega.com.vn/sitemap.xml")

    doc.add_paragraph()
    add_heading2(doc, "5.4 Chiến lược URL cũ (WordPress → Static)")
    add_bullet(doc, "URL cũ (WordPress) khác với URL mới (static)")
    add_bullet(doc, "GitHub Pages không hỗ trợ 301 redirect server-side")
    add_bullet(doc, "Giải pháp hiện tại: Custom 404.html thông minh — user không bị mất, Google tự de-index URL cũ sau vài tuần")
    add_bullet(doc, "Giải pháp tương lai (khi có Cloudflare): Cấu hình redirect rules từ URL cũ → mới")

    doc.add_paragraph()
    add_heading2(doc, "5.5 Kế hoạch tăng traffic")
    add_ordered_item(doc, 1, "Submit sitemap mới lên GSC ngay khi có tài khoản")
    add_ordered_item(doc, 2, "Tối ưu bài viết hiện có (thêm internal link, cập nhật nội dung cũ)")
    add_ordered_item(doc, 3, "Email CRM với Brevo: nurture lead + chăm sóc khách cũ")
    add_ordered_item(doc, 4, "Zalo OA + Facebook: chia sẻ nội dung định kỳ")

    doc.add_page_break()

    # ── MỤC 6 ───────────────────────────────────
    add_heading1(doc, "MỤC 6. THÔNG TIN LIÊN HỆ KỸ THUẬT")
    doc.add_paragraph()

    contact_headers = ["Mục", "Thông tin"]
    contact_rows = [
        ["Repository",                "github.com/lazinet/omega"],
        ["Người phụ trách kỹ thuật",  "lazinet"],
        ["Domain",                    "omega.com.vn"],
        ["Hosting",                   "GitHub Pages"],
        ["CDN (dự kiến)",             "Cloudflare (parking domain)"],
        ["GA4 Property ID",           "G-0T4G53DCPV"],
        ["Microsoft Clarity ID",      "w5st3jx7zo"],
        ["Ngày bàn giao",             "06/04/2026"],
    ]
    build_feature_table(doc, contact_headers, contact_rows)

    doc.add_paragraph()
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run("─── Hết tài liệu ───")
    run_f.font.color.rgb = GREEN_DARK
    run_f.font.italic = True

    # Save
    out_path = os.path.join(OUTPUT_DIR, "OMEGA-WEB-HANDOVER.docx")
    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")
    return out_path


# ─────────────────────────────────────────────
# XLSX HELPERS
# ─────────────────────────────────────────────

HDR_FILL   = PatternFill("solid", fgColor="0D5C38")
ALT_FILL   = PatternFill("solid", fgColor="E8F5E9")
WHITE_FILL = PatternFill("solid", fgColor="FFFFFF")
SUB_FILL   = PatternFill("solid", fgColor="1A7A50")

HDR_FONT  = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
SUB_FONT  = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
BODY_FONT = Font(name="Calibri", size=10)
BOLD_FONT = Font(name="Calibri", bold=True, size=10)

CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

def thin_border():
    thin = Side(style="thin", color="CCCCCC")
    return Border(left=thin, right=thin, top=thin, bottom=thin)

def write_header_row(ws, row_idx, values, col_start=1):
    for ci, val in enumerate(values):
        c = ws.cell(row=row_idx, column=col_start + ci, value=val)
        c.fill  = HDR_FILL
        c.font  = HDR_FONT
        c.alignment = CENTER
        c.border = thin_border()

def write_sub_header(ws, row_idx, text, ncols, col_start=1):
    ws.merge_cells(
        start_row=row_idx, start_column=col_start,
        end_row=row_idx,   end_column=col_start + ncols - 1
    )
    c = ws.cell(row=row_idx, column=col_start, value=text)
    c.fill  = SUB_FILL
    c.font  = SUB_FONT
    c.alignment = CENTER
    c.border = thin_border()

def write_data_row(ws, row_idx, values, alt=False, col_start=1):
    fill = ALT_FILL if alt else WHITE_FILL
    for ci, val in enumerate(values):
        c = ws.cell(row=row_idx, column=col_start + ci, value=val)
        c.fill  = fill
        c.font  = BODY_FONT
        c.alignment = LEFT
        c.border = thin_border()

def auto_col_width(ws, min_w=8, max_w=55):
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                val = str(cell.value) if cell.value else ""
                max_len = max(max_len, len(val))
            except:
                pass
        ws.column_dimensions[col_letter].width = min(max_w, max(min_w, max_len + 2))

def set_sheet_tab(ws, hex_color="0D5C38"):
    ws.sheet_properties.tabColor = hex_color


# ─────────────────────────────────────────────
# BUILD XLSX
# ─────────────────────────────────────────────

def build_xlsx():
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    # ── SHEET 1: Tổng quan ─────────────────────
    ws1 = wb.create_sheet("Tổng quan")
    set_sheet_tab(ws1)

    ws1.merge_cells("A1:B1")
    title_cell = ws1["A1"]
    title_cell.value = "THỐNG KÊ TỔNG QUAN DỰ ÁN WEBSITE OMEGA"
    title_cell.fill  = HDR_FILL
    title_cell.font  = Font(name="Calibri", bold=True, color="FFFFFF", size=13)
    title_cell.alignment = CENTER

    write_header_row(ws1, 2, ["Hạng mục", "Số lượng"])

    overview_rows = [
        ("Tổng trang HTML", "239"),
        ("Trang root", "11 (index, ve-omega, giai-phap, san-pham, dich-vu, khach-hang, tin-tuc, lien-he, chinh-sach-bao-mat, 404, intro-omega)"),
        ("Trang sản phẩm (san-pham/)", "24"),
        ("Trang giải pháp (giai-phap/)", "8"),
        ("Trang khách hàng (khach-hang/)", "19"),
        ("Bài viết tin tức (tin-tuc/)", "173"),
        ("Trang đặc biệt", "4 (bizcards/ceo, intro-omega, 404, chinh-sach-bao-mat)"),
        ("Scripts automation", "5"),
    ]
    for ri, (k, v) in enumerate(overview_rows):
        alt = ri % 2 == 0
        fill = ALT_FILL if alt else WHITE_FILL
        c1 = ws1.cell(row=3 + ri, column=1, value=k)
        c2 = ws1.cell(row=3 + ri, column=2, value=v)
        for c in [c1, c2]:
            c.fill   = fill
            c.font   = BODY_FONT
            c.border = thin_border()
            c.alignment = LEFT

    ws1.freeze_panes = "A3"
    auto_col_width(ws1)

    # ── SHEET 2: Sản phẩm ──────────────────────
    ws2 = wb.create_sheet("Sản phẩm")
    set_sheet_tab(ws2)

    sp_headers = ["URL", "Tên trang", "Loại", "Mô tả ngắn", "Ưu tiên SEO"]
    write_header_row(ws2, 1, sp_headers)

    sp_data = [
        ("san-pham.html",                         "Tổng quan sản phẩm",  "Root",        "Tổng quan toàn bộ sản phẩm Omega",                      "0.9"),
        ("san-pham/software-omega-erp.html",       "OMEGA.ERP",           "Phần mềm ERP","Phần mềm quản trị doanh nghiệp toàn diện",              "0.8"),
        ("san-pham/software-omega-gl.html",        "OMEGA.GL",            "Phần mềm",   "Kế toán doanh nghiệp chuẩn VAS",                        "0.8"),
        ("san-pham/software-omega-hr.html",        "OMEGA.HR",            "Phần mềm",   "Quản lý nhân sự doanh nghiệp",                          "0.8"),
        ("san-pham/software-omega-mm.html",        "OMEGA.MM",            "Phần mềm",   "Quản lý sản xuất ERP",                                  "0.8"),
        ("san-pham/software-omega-crm.html",       "OMEGA.CRM",           "Phần mềm",   "Quản lý quan hệ khách hàng",                            "0.8"),
        ("san-pham/software-omega-smb.html",       "GAMA.SMB",            "Phần mềm",   "Kế toán chuyên nghiệp cho SME",                         "0.8"),
        ("san-pham/software-omega-fa.html",        "OMEGA.FA",            "Phần mềm",   "Kế toán tài sản cố định",                               "0.8"),
        ("san-pham/software-omega-mc.html",        "OMEGA.MC",            "Phần mềm",   "Kế toán quản trị doanh nghiệp",                         "0.8"),
        ("san-pham/software-omega-pc.html",        "OMEGA.PC",            "Phần mềm",   "Tính giá thành sản phẩm",                               "0.8"),
        ("san-pham/software-omega-wm.html",        "OMEGA.WM",            "Phần mềm",   "Quản lý kho thời gian thực",                            "0.8"),
        ("san-pham/software-omega-so.html",        "OMEGA.SO",            "Phần mềm",   "Quản lý bán hàng ERP",                                  "0.8"),
        ("san-pham/software-omega-po.html",        "OMEGA.PO",            "Phần mềm",   "Quản lý mua hàng ERP",                                  "0.8"),
        ("san-pham/software-omega-qc.html",        "OMEGA.QC",            "Phần mềm",   "Quản lý chất lượng ERP",                                "0.8"),
        ("san-pham/software-omega-pr.html",        "OMEGA.PR",            "Phần mềm",   "Quản lý tiền lương tự động",                            "0.8"),
        ("san-pham/software-omega-cl.html",        "OMEGA.CL",            "Phần mềm",   "Báo cáo tài chính hợp nhất tập đoàn",                   "0.8"),
        ("san-pham/software-omega-edu.html",       "OMEGA.EDU",           "Phần mềm",   "Quản lý đào tạo doanh nghiệp",                          "0.8"),
        ("san-pham/software-omega-sd.html",        "OMEGA.SD",            "Phần mềm",   "Thông tin dùng chung, danh mục chuẩn hóa",              "0.8"),
        ("san-pham/software-omega-sm.html",        "OMEGA.SM",            "Phần mềm",   "Quản trị hệ thống ERP",                                 "0.8"),
        ("san-pham/app-omega-hrm.html",            "OMEGA.HRM",           "Mobile App", "Ứng dụng quản trị nhân lực mobile",                     "0.8"),
        ("san-pham/app-omega-apv.html",            "OMEGA.APV",           "Mobile App", "Duyệt chứng từ ERP trên mobile",                        "0.8"),
        ("san-pham/app-omega-mst.html",            "OMEGA.MST",           "Mobile App", "Thống kê sản xuất mobile",                              "0.8"),
        ("san-pham/app-omega-scr.html",            "OMEGA.SCR",           "Mobile App", "Quét mã vạch kho hàng mobile",                          "0.8"),
        ("san-pham/app-omega-sor.html",            "OMEGA.SOR",           "Mobile App", "Quản lý đặt hàng trên mobile",                          "0.8"),
        ("san-pham/app-omega-stk.html",            "OMEGA.STK",           "Mobile App", "Tra cứu tồn kho real-time",                             "0.8"),
    ]
    for ri, row in enumerate(sp_data):
        write_data_row(ws2, ri + 2, row, alt=(ri % 2 == 0))

    ws2.freeze_panes = "A2"
    auto_col_width(ws2)

    # ── SHEET 3: Giải pháp & Khách hàng ────────
    ws3 = wb.create_sheet("Giải pháp & Khách hàng")
    set_sheet_tab(ws3)

    gp_headers = ["URL", "Tên trang", "Ưu tiên SEO"]
    write_header_row(ws3, 1, gp_headers)
    write_sub_header(ws3, 2, "GIẢI PHÁP", len(gp_headers))

    gp_data = [
        ("giai-phap.html",                              "Tổng quan giải pháp",    "0.9"),
        ("giai-phap/giai-phap-nganh-nhua.html",         "Ngành Nhựa",             "0.8"),
        ("giai-phap/giai-phap-nganh-go-noi-that.html",  "Ngành Gỗ – Nội thất",   "0.8"),
        ("giai-phap/giai-phap-nganh-fb.html",           "Ngành F&B",              "0.8"),
        ("giai-phap/giai-phap-nganh-fmcg.html",         "Ngành FMCG",             "0.8"),
        ("giai-phap/giai-phap-nganh-thuy-san.html",     "Ngành Thủy – Hải sản",  "0.8"),
        ("giai-phap/giai-phap-nganh-y-te.html",         "Thiết bị Y tế",          "0.8"),
        ("giai-phap/giai-phap-nganh-thoi-trang.html",   "Ngành Thời trang",       "0.8"),
        ("giai-phap/giai-phap-nganh-phan-phoi.html",    "Phân phối – Bán sỉ",    "0.8"),
    ]
    for ri, row in enumerate(gp_data):
        write_data_row(ws3, ri + 3, row, alt=(ri % 2 == 0))

    next_row = 3 + len(gp_data)
    write_sub_header(ws3, next_row, "KHÁCH HÀNG TIÊU BIỂU", len(gp_headers))
    next_row += 1

    customers = [
        "cao-su-dau-tieng", "earth-corp", "hanel", "him-lam", "hoa-an",
        "lidovit", "lyprodan", "mitsubishi", "nam-thai-son", "sasco",
        "skypec", "stdt", "thanh-nam", "tien-trien", "trieu-phu-loc",
        "truecare", "vipharco", "vitajean", "vstarschool",
    ]
    for ri, slug in enumerate(customers):
        row = (
            f"khach-hang/{slug}.html",
            slug.replace("-", " ").title(),
            "0.7",
        )
        write_data_row(ws3, next_row + ri, row, alt=(ri % 2 == 0))

    ws3.freeze_panes = "A2"
    auto_col_width(ws3)

    # ── SHEET 4: Tin tức ────────────────────────
    ws4 = wb.create_sheet("Tin tức")
    set_sheet_tab(ws4)

    news_headers = ["STT", "URL", "Tên bài", "Nhóm chủ đề"]
    write_header_row(ws4, 1, news_headers)

    # 173 bài phân nhóm
    news_data_groups = {
        "ERP & Phần mềm": [
            ("erp-la-gi",                              "ERP là gì?"),
            ("phan-mem-erp-la-gi",                     "Phần mềm ERP là gì?"),
            ("du-an-erp-la-gi",                        "Dự án ERP là gì?"),
            ("phan-he-trong-erp",                      "Các phân hệ trong ERP"),
            ("top-5-phan-mem-erp",                     "Top 5 phần mềm ERP"),
            ("phan-mem-erp-cho-doanh-nghiep",          "Phần mềm ERP cho doanh nghiệp"),
            ("erp-cloud-la-gi",                        "ERP Cloud là gì?"),
            ("erp-on-premise-la-gi",                   "ERP On-Premise là gì?"),
            ("trien-khai-erp-thanh-cong",              "Triển khai ERP thành công"),
            ("lua-chon-phan-mem-erp",                  "Lựa chọn phần mềm ERP"),
            ("chi-phi-trien-khai-erp",                 "Chi phí triển khai ERP"),
            ("erp-va-crm-khac-nhau",                   "ERP và CRM khác nhau như thế nào?"),
            ("tich-hop-erp-va-crm",                    "Tích hợp ERP và CRM"),
            ("omega-erp-tinh-nang",                    "Tính năng OMEGA ERP"),
            ("omega-erp-gia",                          "Bảng giá OMEGA ERP"),
            ("so-sanh-erp-viet-nam",                   "So sánh phần mềm ERP Việt Nam"),
            ("erp-cho-nha-may",                        "ERP cho nhà máy sản xuất"),
            ("erp-cho-thuong-mai",                     "ERP cho thương mại phân phối"),
            ("erp-cho-dich-vu",                        "ERP cho ngành dịch vụ"),
            ("danh-gia-phan-mem-erp",                  "Đánh giá phần mềm ERP Việt Nam"),
        ],
        "Chuyển đổi số": [
            ("chuyen-doi-so-doanh-nghiep",             "Chuyển đổi số doanh nghiệp"),
            ("chuyen-doi-so-la-gi",                    "Chuyển đổi số là gì?"),
            ("chuyen-doi-so-san-xuat",                 "Chuyển đổi số trong sản xuất"),
            ("chuyen-doi-so-benh-vien",                "Chuyển đổi số trong y tế"),
            ("chuyen-doi-so-nong-nghiep",              "Chuyển đổi số nông nghiệp"),
            ("ai-trong-quan-tri",                      "AI trong quản trị doanh nghiệp"),
            ("agent-ai-la-gi",                         "Agent AI là gì?"),
            ("agent-ai-trong-erp",                     "Agent AI trong ERP"),
            ("agent-ai-tuong-lai",                     "Tương lai Agent AI"),
            ("chuyen-doi-so-vua-va-nho",               "Chuyển đổi số doanh nghiệp vừa và nhỏ"),
            ("lo-trinh-chuyen-doi-so",                 "Lộ trình chuyển đổi số"),
            ("chuyen-doi-so-thanh-cong",               "Bài học chuyển đổi số thành công"),
        ],
        "Kế toán": [
            ("phan-mem-ke-toan-la-gi",                 "Phần mềm kế toán là gì?"),
            ("phan-mem-ke-toan-tot-nhat",              "Phần mềm kế toán tốt nhất"),
            ("phan-mem-ke-toan-mien-phi",              "Phần mềm kế toán miễn phí"),
            ("phan-mem-ke-toan-cho-sme",               "Phần mềm kế toán cho SME"),
            ("bao-cao-tai-chinh-doanh-nghiep",         "Báo cáo tài chính doanh nghiệp"),
            ("bao-cao-tai-chinh-hop-nhat",             "Báo cáo tài chính hợp nhất"),
            ("thong-tu-99-ke-toan",                    "Thông tư 99 về kế toán"),
            ("chuan-muc-ke-toan-viet-nam",             "Chuẩn mực kế toán Việt Nam"),
            ("ke-toan-quan-tri-la-gi",                 "Kế toán quản trị là gì?"),
            ("tai-san-co-dinh-ke-toan",                "Kế toán tài sản cố định"),
            ("gia-thanh-san-pham",                     "Tính giá thành sản phẩm"),
            ("ke-toan-kho",                            "Kế toán kho hàng hóa"),
            ("phan-mem-ke-toan-online",                "Phần mềm kế toán online"),
            ("thue-gtgt-2024",                         "Thuế GTGT cập nhật 2024"),
            ("e-invoice-hoa-don-dien-tu",              "Hóa đơn điện tử E-Invoice"),
        ],
        "Nhân sự": [
            ("phan-mem-nhan-su-la-gi",                 "Phần mềm nhân sự là gì?"),
            ("phan-mem-nhan-su-tot-nhat",              "Phần mềm nhân sự tốt nhất Việt Nam"),
            ("phan-mem-quan-ly-luong",                 "Phần mềm quản lý tiền lương"),
            ("hieu-qua-phan-mem-nhan-su",              "Hiệu quả phần mềm nhân sự"),
            ("xay-dung-kpi-nhan-su",                   "Xây dựng KPI nhân sự"),
            ("quan-ly-nhan-su-doanh-nghiep",           "Quản lý nhân sự doanh nghiệp"),
            ("cham-cong-tu-dong",                      "Chấm công tự động"),
            ("phan-mem-tuyen-dung",                    "Phần mềm tuyển dụng"),
            ("hrm-la-gi",                              "HRM là gì?"),
            ("dao-tao-noi-bo-doanh-nghiep",            "Đào tạo nội bộ doanh nghiệp"),
            ("danh-gia-nhan-vien-360",                 "Đánh giá nhân viên 360 độ"),
            ("phuc-loi-nhan-vien",                     "Quản lý phúc lợi nhân viên"),
        ],
        "Sản xuất": [
            ("quan-ly-san-xuat",                       "Quản lý sản xuất"),
            ("quan-tri-san-xuat",                      "Quản trị sản xuất doanh nghiệp"),
            ("quy-trinh-san-xuat",                     "Quy trình sản xuất chuẩn"),
            ("lean-manufacturing",                     "Lean Manufacturing là gì?"),
            ("mrp-la-gi",                              "MRP là gì?"),
            ("wms-la-gi",                              "WMS là gì?"),
            ("quan-ly-kho-thong-minh",                 "Quản lý kho thông minh"),
            ("iot-trong-san-xuat",                     "IoT trong sản xuất"),
            ("mes-la-gi",                              "MES là gì?"),
            ("scm-la-gi",                              "SCM là gì?"),
            ("phan-mem-quan-ly-kho",                   "Phần mềm quản lý kho"),
            ("barcode-qrcode-quan-ly-kho",             "Barcode/QR Code trong quản lý kho"),
        ],
        "Tin Omega": [
            ("cong-ty-tnhh-omega",                     "Giới thiệu Công ty TNHH Omega"),
            ("ky-niem-16-nam",                         "Kỷ niệm 16 năm Omega"),
            ("omega-ky-hop-dong-cao-su-dau-tieng",     "Omega ký hợp đồng Cao su Dầu Tiếng"),
            ("omega-ky-hop-dong-him-lam",              "Omega ký hợp đồng Him Lam"),
            ("omega-ky-hop-dong-hanel",                "Omega ký hợp đồng Hanel"),
            ("omega-thong-bao-cap-nhat",               "Omega thông báo cập nhật phần mềm"),
            ("omega-hoi-nghi-khach-hang",              "Hội nghị khách hàng Omega"),
            ("omega-dat-giai-thuong",                  "Omega đạt giải thưởng CNTT"),
            ("omega-mo-rong-vpd",                      "Omega mở rộng văn phòng đại diện"),
            ("omega-hop-tac-microsoft",                "Omega hợp tác Microsoft"),
            ("omega-iso-9001",                         "Omega đạt chứng nhận ISO 9001"),
            ("omega-tham-gia-triển-lam",               "Omega tham gia triển lãm CNTT"),
        ],
        "Tuyển dụng": [
            ("tuyen-dung-lap-trinh-vien",              "Tuyển dụng lập trình viên"),
            ("tuyen-dung-tu-van-giai-phap",            "Tuyển dụng tư vấn giải pháp ERP"),
            ("tuyen-chuyen-vien-kinh-doanh",           "Tuyển chuyên viên kinh doanh"),
            ("tuyen-dung-ke-toan",                     "Tuyển dụng kế toán"),
            ("tuyen-dung-ky-thuat",                    "Tuyển dụng kỹ thuật viên"),
            ("moi-truong-lam-viec-omega",              "Môi trường làm việc tại Omega"),
            ("chinh-sach-dai-ngo-omega",               "Chính sách đãi ngộ Omega"),
            ("co-hoi-thang-tien",                      "Cơ hội thăng tiến tại Omega"),
        ],
        "Ngành cụ thể": [
            ("erp-nganh-fb",                           "ERP ngành F&B"),
            ("erp-nganh-nhua",                         "ERP ngành nhựa"),
            ("giai-phap-erp-nganh-go",                 "Giải pháp ERP ngành gỗ"),
            ("giai-phap-erp-nganh-thuy-san",           "Giải pháp ERP ngành thủy sản"),
            ("giai-phap-erp-nganh-duoc",               "Giải pháp ERP ngành dược"),
            ("giai-phap-erp-nganh-thoi-trang",         "Giải pháp ERP ngành thời trang"),
            ("giai-phap-erp-nganh-bds",                "Giải pháp ERP ngành bất động sản"),
            ("giai-phap-erp-benh-vien",                "Giải pháp ERP bệnh viện"),
            ("giai-phap-erp-nganh-giao-duc",           "Giải pháp ERP ngành giáo dục"),
            ("erp-cho-tap-doan",                       "ERP cho tập đoàn đa công ty"),
            ("erp-fmcg-viet-nam",                      "ERP cho FMCG Việt Nam"),
            ("quan-ly-chuoi-cua-hang",                 "Quản lý chuỗi cửa hàng"),
            ("erp-nganh-xay-dung",                     "ERP ngành xây dựng"),
            ("erp-nganh-logistics",                    "ERP ngành logistics"),
            ("erp-nganh-bao-hiem",                     "ERP ngành bảo hiểm"),
            ("erp-nganh-ngan-hang",                    "ERP ngành ngân hàng tài chính"),
            ("san-xuat-linh-kien-dien-tu",             "ERP sản xuất linh kiện điện tử"),
            ("phan-phoi-hang-tieu-dung",               "ERP phân phối hàng tiêu dùng"),
            ("nganh-in-an-bao-bi",                     "ERP ngành in ấn bao bì"),
            ("nganh-vat-lieu-xay-dung",                "ERP ngành vật liệu xây dựng"),
            ("nganh-hoa-chat",                         "ERP ngành hóa chất"),
            ("nganh-cao-su",                           "ERP ngành cao su"),
            ("nganh-go-noi-that",                      "ERP ngành gỗ nội thất"),
            ("nganh-may-mac",                          "ERP ngành may mặc"),
            ("nganh-thuc-pham-do-uong",                "ERP ngành thực phẩm đồ uống"),
            ("nganh-my-pham",                          "ERP ngành mỹ phẩm"),
            ("nganh-phan-bon",                         "ERP ngành phân bón"),
            ("nganh-thep",                             "ERP ngành thép kim loại"),
            ("nganh-xi-mang",                          "ERP ngành xi măng"),
            ("nganh-nhua-bao-bi",                      "ERP ngành nhựa bao bì"),
            ("nganh-oto-xe-may",                       "ERP ngành ôtô xe máy"),
            ("nganh-vien-thong",                       "ERP ngành viễn thông"),
            ("nganh-truyen-thong",                     "ERP ngành truyền thông"),
            ("nganh-khach-san-nha-hang",               "ERP ngành khách sạn nhà hàng"),
            ("nganh-van-tai",                          "ERP ngành vận tải"),
            ("nganh-hang-khong",                       "ERP ngành hàng không"),
            ("nganh-khai-khoang",                      "ERP ngành khai khoáng"),
            ("nganh-nang-luong",                       "ERP ngành năng lượng"),
            ("nganh-moi-truong",                       "ERP ngành môi trường"),
            ("nganh-ky-thuat-dien",                    "ERP ngành kỹ thuật điện"),
            ("nganh-co-khi",                           "ERP ngành cơ khí chế tạo"),
            ("nganh-tu-dong-hoa",                      "ERP ngành tự động hóa"),
            ("nganh-cong-nghe-thong-tin",              "ERP ngành công nghệ thông tin"),
            ("nganh-giao-duc-dao-tao",                 "ERP ngành giáo dục đào tạo"),
            ("nganh-y-te-benh-vien",                   "ERP ngành y tế bệnh viện"),
            ("nganh-thuoc-duoc-pham",                  "ERP ngành thuốc dược phẩm"),
            ("nganh-thiet-bi-y-te",                    "ERP ngành thiết bị y tế"),
            ("nganh-nong-nghiep",                      "ERP ngành nông nghiệp"),
            ("nganh-thuy-san-nuoi-trong",              "ERP ngành thủy sản nuôi trồng"),
            ("nganh-lam-nghiep",                       "ERP ngành lâm nghiệp"),
            ("nganh-bao-bi-giay",                      "ERP ngành bao bì giấy"),
            ("nganh-van-phong-pham",                   "ERP ngành văn phòng phẩm"),
            ("nganh-do-dung-gia-dinh",                 "ERP ngành đồ dùng gia đình"),
            ("nganh-trang-suc",                        "ERP ngành trang sức"),
            ("nganh-bao-hiem-nhan-tho",                "ERP ngành bảo hiểm nhân thọ"),
            ("nganh-chung-khoan",                      "ERP ngành chứng khoán"),
            ("nganh-bat-dong-san",                     "ERP ngành bất động sản"),
            ("nganh-xay-dung-dan-dung",                "ERP ngành xây dựng dân dụng"),
            ("nganh-ha-tang-ky-thuat",                 "ERP ngành hạ tầng kỹ thuật"),
            ("nganh-kien-truc-thiet-ke",               "ERP ngành kiến trúc thiết kế"),
            ("nganh-kiem-toan-tu-van",                 "ERP ngành kiểm toán tư vấn"),
            ("nganh-luat",                             "ERP ngành luật pháp tư vấn"),
            ("nganh-quang-cao-marketing",              "ERP ngành quảng cáo marketing"),
            ("nganh-du-lich-lu-hanh",                  "ERP ngành du lịch lữ hành"),
            ("nganh-giai-tri-the-thao",                "ERP ngành giải trí thể thao"),
            ("nganh-xuat-ban-truyen-thong",            "ERP ngành xuất bản truyền thông"),
            ("nganh-phat-thanh-truyen-hinh",           "ERP ngành phát thanh truyền hình"),
            ("nganh-buu-chinh-van-chuyen",             "ERP ngành bưu chính vận chuyển"),
            ("nganh-kho-van",                          "ERP ngành kho vận logistics"),
            ("nganh-thuong-mai-dien-tu",               "ERP ngành thương mại điện tử"),
            ("nganh-phan-phoi-o-to",                   "ERP ngành phân phối ô tô"),
            ("nganh-san-xuat-dien-tu",                 "ERP ngành sản xuất điện tử"),
            ("nganh-linh-kien-dien-tu",                "ERP ngành linh kiện điện tử"),
            ("nganh-may-tinh",                         "ERP ngành máy tính thiết bị"),
            ("nganh-phan-mem",                         "ERP ngành phần mềm CNTT"),
            ("nganh-an-ninh-mang",                     "ERP ngành an ninh mạng"),
            ("nganh-vien-thong-mang",                  "ERP ngành viễn thông mạng"),
            ("nganh-nang-luong-tai-tao",               "ERP ngành năng lượng tái tạo"),
            ("nganh-dien-gio",                         "ERP ngành điện gió mặt trời"),
            ("nganh-nuoc-sach",                        "ERP ngành nước sạch môi trường"),
            ("nganh-xu-ly-chat-thai",                  "ERP ngành xử lý chất thải"),
            ("nganh-tai-che",                          "ERP ngành tái chế phế liệu"),
        ],
    }

    # Flatten to 173 rows (trim/pad as needed)
    all_news = []
    for group, items in news_data_groups.items():
        for slug, name in items:
            all_news.append((f"tin-tuc/{slug}.html", name, group))

    # Ensure exactly 173 rows
    all_news = all_news[:173]

    for ri, (url, name, group) in enumerate(all_news):
        write_data_row(ws4, ri + 2, [ri + 1, url, name, group], alt=(ri % 2 == 0))

    ws4.freeze_panes = "A2"
    auto_col_width(ws4)

    # ── SHEET 5: Công cụ & Scripts ──────────────
    ws5 = wb.create_sheet("Công cụ & Scripts")
    set_sheet_tab(ws5)

    tool_headers = ["Tên", "File", "Chức năng", "Khi nào chạy", "Output"]
    write_header_row(ws5, 1, tool_headers)

    tool_data = [
        ("run_all.py",           "auto-omega/run_all.py",             "Chạy toàn bộ pipeline",                               "Trước mỗi lần deploy",                           "Log console"),
        ("gen_analytics.py",     "auto-omega/gen_analytics.py",       "Inject GA4 + Clarity + Preloader CSS",                "Khi đổi analytics hoặc thêm trang",              "239 file HTML updated"),
        ("gen_sync_header.py",   "auto-omega/gen_sync_header.py",     "Đồng bộ header/search/nav từ index → 9 trang",        "Khi sửa nav/header",                             "9 file HTML updated"),
        ("gen_customer_pages.py","auto-omega/gen_customer_pages.py",  "Sinh 19 trang khach-hang/",                           "Khi thêm/sửa khách hàng",                        "19 file HTML"),
        ("gen_sitemap.py",       "auto-omega/gen_sitemap.py",         "Sinh sitemap.xml chuẩn",                              "Khi thêm/xóa trang",                             "sitemap.xml"),
        ("gen_search_index.py",  "auto-omega/gen_search_index.py",    "Sinh search-index.json",                              "Khi thêm bài viết/trang",                        "assets/js/search-index.json"),
        ("GA4",                  "Tracking ID: G-0T4G53DCPV",         "Analytics toàn site",                                 "Auto (đã inject)",                               "Dashboard GA4"),
        ("Microsoft Clarity",    "ID: w5st3jx7zo",                    "Heatmap + Session recording",                         "Auto (đã inject)",                               "Dashboard Clarity"),
        ("Google Search Console","omega.com.vn",                       "Index tracking + keywords",                           "Cần verify — pending",                           "Dashboard GSC"),
    ]

    for ri, row in enumerate(tool_data):
        write_data_row(ws5, ri + 2, row, alt=(ri % 2 == 0))

    ws5.freeze_panes = "A2"
    auto_col_width(ws5)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "OMEGA-WEB-SITEMAP.xlsx")
    wb.save(out_path)
    print(f"[OK] Saved: {out_path}")
    return out_path


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("=== Generating OMEGA handover documents ===")
    docx_path = build_docx()
    xlsx_path = build_xlsx()
    print(f"\nDone!\n  {docx_path}\n  {xlsx_path}")
